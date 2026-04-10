"""
parser.py — Extract raw text from PDF or DOCX files.
"""
import io
from pathlib import Path


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Return plain text from a PDF or DOCX upload."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(file_bytes: bytes) -> str:
    import fitz  # pymupdf
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    return "\n".join(pages)


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also pull table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
